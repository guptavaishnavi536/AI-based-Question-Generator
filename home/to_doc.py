from docx import Document

def todoc(content, string):
    # Create a new Document
    doc = Document()

    # Add a paragraph to the document
    paragraph = doc.add_paragraph()

    # Add text to the paragraph
    paragraph.add_run(content)

    # Save the document to a file
    doc.save(f"media/pdf/{string}.docx")
